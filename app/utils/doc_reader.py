from pathlib import Path

from docx import Document


class DocReader:

    def read(self, file_path: Path) -> str:
        document = Document(str(file_path))

        extracted_text = []

        # Read paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                extracted_text.append(text)

        # Read tables
        for table in document.tables:
            for row in table.rows:

                values = []

                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        values.append(text)

                if values:
                    extracted_text.append(
                        " | ".join(values)
                    )

        return "\n".join(extracted_text)
    